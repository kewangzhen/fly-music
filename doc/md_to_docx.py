#!/usr/bin/env python3
"""
Markdown 转 DOCX 转换器
完整保留标题、段落和表格
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

def is_table_line(line):
    """判断是否是表格行"""
    stripped = line.strip()
    return stripped.startswith('|') and stripped.endswith('|')

def is_table_separator(line):
    """判断是否是表格分隔行"""
    stripped = line.strip()
    cells = [c.strip() for c in stripped.split('|')[1:-1]]
    return len(cells) > 0 and all(c == '' or set(c.replace('-', '').replace(':', '')) <= {''} for c in cells)

def parse_markdown_file(md_file):
    """解析markdown文件内容"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    sections = []
    lines = content.split('\n')
    
    i = 0
    current_section = None
    current_content = []
    
    while i < len(lines):
        line = lines[i]
        
        # 标题行
        if line.startswith('#'):
            # 保存之前的内容
            if current_section:
                sections.append((current_section, current_content))
            
            # 提取标题级别和文本
            match = re.match(r'^(#+)\s+(.*)', line)
            if match:
                level = len(match.group(1))
                text = match.group(2)
                current_section = ('title', level, text)
                current_content = []
        
        # 表格行
        elif is_table_line(line):
            table_lines = []
            while i < len(lines) and is_table_line(lines[i]):
                if not is_table_separator(lines[i]):
                    table_lines.append(lines[i])
                i += 1
            # 保存表格
            if table_lines:
                sections.append(('table', table_lines))
            continue
        
        else:
            if line.strip():
                current_content.append(line)
        
        i += 1
    
    # 保存最后的内容
    if current_section:
        sections.append((current_section, current_content))
    
    return sections

def create_docx(md_file, docx_file):
    """创建docx文件"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    
    sections = parse_markdown_file(md_file)
    
    for section in sections:
        if section[0] == 'title':
            level = section[1]
            text = section[2]
            if level == 1:
                doc.add_heading(text, 0)
            elif level == 2:
                doc.add_heading(text, 1)
            elif level == 3:
                doc.add_heading(text, 2)
            else:
                doc.add_heading(text, 3)
        
        elif section[0] == 'paragraph':
            for line in section[1]:
                if line.strip():
                    doc.add_paragraph(line.strip())
        
        elif section[0] == 'table':
            table_lines = section[1]
            if not table_lines:
                continue
            
            # 解析表头
            headers = [h.strip() for h in table_lines[0].split('|')[1:-1]]
            
            # 创建表格
            table = doc.add_table(rows=len(table_lines), cols=len(headers))
            table.style = 'Table Grid'
            
            # 填充数据
            for row_idx, row_line in enumerate(table_lines):
                cells = [c.strip() for c in row_line.split('|')[1:-1]]
                for col_idx, cell_text in enumerate(cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = cell_text
                    
                    # 表头加粗
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
            
            doc.add_paragraph('')
        
        elif section[0] == 'content':
            for line in section[1]:
                if line.strip():
                    doc.add_paragraph(line.strip())
    
    doc.save(docx_file)
    print(f"转换成功: {docx_file}")

if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(script_dir, '测试用例.md')
    docx_file = os.path.join(script_dir, '测试用例.docx')
    
    if not os.path.exists(md_file):
        print(f"错误: 文件不存在 {md_file}")
        exit(1)
    
    create_docx(md_file, docx_file)
